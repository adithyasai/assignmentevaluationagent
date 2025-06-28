"""
Sample data generator for testing the Assignment Agent.
"""
import pandas as pd
from pathlib import Path
from docx import Document
import config

def create_sample_excel():
    """Create a sample Excel file with student data."""
    sample_data = {
        'Name': [
            'Alice Johnson',
            'Bob Smith',
            'Carol Davis',
            'David Wilson',
            'Eva Garcia',
            'Frank Brown',
            'Grace Lee',
            'Henry Taylor'
        ],
        'StudentID': [
            'CS2024001',
            'CS2024002', 
            'CS2024003',
            'CS2024004',
            'CS2024005',
            'CS2024006',
            'CS2024007',
            'CS2024008'
        ],
        'Email': [
            'alice.johnson@university.edu',
            'bob.smith@university.edu',
            'carol.davis@university.edu',
            'david.wilson@university.edu',
            'eva.garcia@university.edu',
            'frank.brown@university.edu',
            'grace.lee@university.edu',
            'henry.taylor@university.edu'
        ],
        'GitHubRepoURL': [
            'https://github.com/facebook/create-react-app',  # Known working React repo
            'https://github.com/vercel/next.js',             # Another working repo
            'https://github.com/microsoft/vscode',           # Non-React repo (for testing)
            'https://github.com/invalid-user/repo-404',      # Invalid repo (for testing)
            'https://github.com/reactjs/react.dev',          # React docs repo
            'https://github.com/vitejs/vite',                # Vite repo
            'https://github.com/remix-run/remix',            # Remix repo
            'https://github.com/storybookjs/storybook'       # Storybook repo
        ]
    }
    
    df = pd.DataFrame(sample_data)
    output_path = config.DATA_DIR / 'sample_students.xlsx'
    df.to_excel(output_path, index=False)
    
    print(f"Sample Excel file created: {output_path}")
    return str(output_path)

def create_sample_word_doc():
    """Create a sample Word document with requirements."""
    doc = Document()
    
    # Title
    title = doc.add_heading('React Assignment Requirements', 0)
    
    # Technical Requirements
    doc.add_heading('Technical Requirements (40 points)', level=1)
    tech_requirements = [
        'Project must build successfully without errors',
        'Use Create React App or Vite as build tool',
        'Include package.json with proper dependencies',
        'Application must start with npm start or npm run dev',
        'No critical console errors in browser'
    ]
    
    for req in tech_requirements:
        doc.add_paragraph(req, style='List Bullet')
    
    # Component Structure
    doc.add_heading('Component Structure (30 points)', level=1)
    component_requirements = [
        'Minimum 3 functional components',
        'Proper component composition and hierarchy',
        'Use of props for data passing between components',
        'State management with useState hook',
        'Event handling implementation'
    ]
    
    for req in component_requirements:
        doc.add_paragraph(req, style='List Bullet')
    
    # Styling & UI
    doc.add_heading('Styling & UI (20 points)', level=1)
    ui_requirements = [
        'Responsive design implementation',
        'CSS modules, styled-components, or CSS-in-JS usage',
        'Clean and professional appearance',
        'Mobile-friendly interface',
        'Consistent styling across components'
    ]
    
    for req in ui_requirements:
        doc.add_paragraph(req, style='List Bullet')
    
    # Code Quality
    doc.add_heading('Code Quality (10 points)', level=1)
    quality_requirements = [
        'Proper file organization and folder structure',
        'Meaningful variable and function names',
        'Clean code practices and readability',
        'Proper commenting where necessary',
        'No unused imports or variables'
    ]
    
    for req in quality_requirements:
        doc.add_paragraph(req, style='List Bullet')
    
    # Additional Notes
    doc.add_heading('Additional Notes', level=1)
    notes = [
        'All repositories should be public on GitHub',
        'Include a README.md with setup instructions',
        'Projects should demonstrate understanding of React fundamentals',
        'Bonus points for creative implementation and extra features'
    ]
    
    for note in notes:
        doc.add_paragraph(note, style='List Bullet')
    
    # Save document
    output_path = config.DATA_DIR / 'sample_requirements.docx'
    doc.save(output_path)
    
    print(f"Sample Word document created: {output_path}")
    return str(output_path)

if __name__ == "__main__":
    # Ensure data directory exists
    config.DATA_DIR.mkdir(exist_ok=True)
    
    # Create sample files
    excel_path = create_sample_excel()
    word_path = create_sample_word_doc()
    
    print(f"\nSample files created:")
    print(f"Excel: {excel_path}")
    print(f"Word: {word_path}")
    print(f"\nYou can use these files to test the Assignment Agent!")
